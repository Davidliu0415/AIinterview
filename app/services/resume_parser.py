from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class ResumeParseError(ValueError):
    """Raised when an uploaded resume cannot be parsed."""


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _parse_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_resume_file(filename: str, content: bytes) -> tuple[str, str]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        text = _parse_pdf(content)
        file_type = "pdf"
    elif suffix == ".docx":
        text = _parse_docx(content)
        file_type = "docx"
    else:
        raise ResumeParseError("仅支持 PDF 和 DOCX 简历。")

    text = normalize_text(text)
    if len(text) < 20:
        raise ResumeParseError("简历文本过短，可能是扫描件或无法解析的文件。")
    return text, file_type
