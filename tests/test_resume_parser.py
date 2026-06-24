from io import BytesIO

import pytest
from docx import Document

from app.services.resume_parser import ResumeParseError, normalize_text, parse_resume_file


def test_normalize_text_compacts_spaces_and_blank_lines():
    assert normalize_text("Java   Spring\n\n\nMySQL") == "Java Spring\n\nMySQL"


def test_parse_docx_resume():
    document = Document()
    document.add_paragraph("张三 Java 后端开发 Spring Boot MySQL Redis 项目经验")
    document.add_paragraph("负责接口开发、数据库设计和基础性能优化")
    buffer = BytesIO()
    document.save(buffer)

    text, file_type = parse_resume_file("resume.docx", buffer.getvalue())

    assert file_type == "docx"
    assert "Spring Boot" in text
    assert "数据库设计" in text


def test_parse_resume_rejects_unsupported_file():
    with pytest.raises(ResumeParseError):
        parse_resume_file("resume.txt", b"Java developer")
